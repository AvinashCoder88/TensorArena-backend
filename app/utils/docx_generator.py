from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

class DocxGenerator:
    def generate_report(self, grading_result: dict, student_name: str = "Student") -> io.BytesIO:
        document = Document()
        
        # Title
        title = document.add_heading(f'Grading Report for {student_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Overall Grade
        grade = grading_result.get("grade", "N/A")
        p = document.add_paragraph()
        run = p.add_run(f"Overall Grade: {grade}")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF) # Blue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Remarks
        document.add_heading('General Remarks', level=1)
        remarks = grading_result.get("remarks", "No remarks provided.")
        document.add_paragraph(remarks)
        
        # Spelling Mistakes
        if grading_result.get("spelling_mistakes"):
            document.add_heading('Spelling Corrections', level=1)
            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Original'
            hdr_cells[1].text = 'Correction'
            hdr_cells[2].text = 'Context'
            
            for mistake in grading_result["spelling_mistakes"]:
                row_cells = table.add_row().cells
                row_cells[0].text = mistake.get("original", "")
                row_cells[1].text = mistake.get("correction", "")
                row_cells[2].text = mistake.get("context", "")
                
        # Math Corrections
        if grading_result.get("math_corrections"):
            document.add_heading('Math Corrections', level=1)
            for item in grading_result["math_corrections"]:
                p = document.add_paragraph()
                p.add_run("Equation: ").bold = True
                p.add_run(item.get("original", ""))
                
                p2 = document.add_paragraph()
                p2.add_run("Correction: ").bold = True
                p2.add_run(item.get("correction", "")).font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Green
                
                p3 = document.add_paragraph()
                p3.add_run("Explanation: ").italic = True
                p3.add_run(item.get("explanation", ""))
                document.add_paragraph("") # Spacer

        # Diagram Analysis
        if grading_result.get("diagram_analysis"):
            document.add_heading('Diagram Analysis', level=1)
            for item in grading_result["diagram_analysis"]:
                p = document.add_paragraph()
                status = "Correct" if item.get("is_correct") else "Incorrect"
                run = p.add_run(f"Diagram: {status}")
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) if item.get("is_correct") else RGBColor(0xFF, 0x00, 0x00)
                
                document.add_paragraph(f"Description: {item.get('description', '')}")
                document.add_paragraph(f"Feedback: {item.get('feedback', '')}")
                document.add_paragraph("") # Spacer

        # Save to bytes
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream
